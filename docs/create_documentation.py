#!/usr/bin/env python3
"""
VibeLink Ghana Technical Documentation Generator
Creates a comprehensive Word document with professional formatting
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def add_cover_page(doc):
    """Add professional cover page"""
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('VibeLink Ghana\n')
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(26, 35, 126)  # Navy blue

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Technical Documentation\n\n')
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.color.rgb = RGBColor(63, 81, 181)

    # Description
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run = desc.add_run('Digital Event Invitation Platform\n\n\n\n')
    desc_run.font.size = Pt(18)
    desc_run.font.italic = True

    # Version info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = info.add_run(f'Version 1.0\n')
    info_text.font.size = Pt(14)

    date_text = info.add_run(f'January 28, 2026\n\n')
    date_text.font.size = Pt(14)

    logo_placeholder = info.add_run('[Logo Placeholder]\n\n')
    logo_placeholder.font.size = Pt(12)
    logo_placeholder.font.italic = True

    website = info.add_run('https://vibelinkgh.com/')
    website.font.size = Pt(12)
    website.font.color.rgb = RGBColor(33, 150, 243)

    doc.add_page_break()

def add_table_of_contents(doc):
    """Add table of contents placeholder"""
    heading = doc.add_heading('TABLE OF CONTENTS', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    toc = doc.add_paragraph()
    toc.add_run('[Table of Contents - Auto-generated in Word]\n\n')
    toc.paragraph_format.line_spacing = 1.15

    note = doc.add_paragraph()
    note.add_run('Note: In Microsoft Word, place cursor here and go to References > Table of Contents > Automatic Table 1')
    note.runs[0].font.italic = True
    note.runs[0].font.size = Pt(10)

    doc.add_page_break()

def add_executive_summary(doc):
    """Add executive summary section"""
    doc.add_heading('EXECUTIVE SUMMARY', level=1)

    doc.add_paragraph(
        'VibeLink Ghana is a comprehensive digital event invitation platform designed specifically '
        'for the Ghanaian market. The platform revolutionizes how Ghanaians create, manage, and '
        'distribute invitations for various events including weddings, funerals, naming ceremonies, '
        'graduations, and corporate events.'
    )

    doc.add_heading('Key Achievements', level=2)
    achievements = [
        'Fully functional digital invitation creation and management system',
        'Integrated payment processing with Paystack for seamless transactions',
        'Comprehensive admin dashboard with real-time analytics',
        'Customer portal for order tracking and management',
        'AI-powered chatbot using Google Gemini for customer support',
        'Multi-channel communication via email and WhatsApp integration',
        'SEO-optimized with Progressive Web App capabilities',
        'Referral program to drive organic growth',
        'Blog system for content marketing',
        'Professional portfolio showcase'
    ]

    for achievement in achievements:
        p = doc.add_paragraph(achievement, style='List Bullet')
        p.paragraph_format.line_spacing = 1.15

    doc.add_heading('Technology Highlights', level=2)
    doc.add_paragraph(
        'The platform leverages modern web technologies including React 18.3.1 with TypeScript '
        'for type safety, Vite for blazing-fast builds, Tailwind CSS with shadcn-ui for beautiful '
        'and accessible UI components, and Supabase as a complete backend-as-a-service solution '
        'providing PostgreSQL database, authentication, and serverless functions.'
    )

    doc.add_paragraph(
        'The architecture follows a modern three-tier deployment model: development on local PC, '
        'version control via GitHub, and production deployment on a dedicated server at '
        '/var/www/vibelinkgh.com/.'
    )

    doc.add_page_break()

def add_project_overview(doc):
    """Add project overview section"""
    doc.add_heading('PROJECT OVERVIEW', level=1)

    doc.add_heading('Business Purpose', level=2)
    doc.add_paragraph(
        'VibeLink Ghana addresses a critical need in the Ghanaian event planning ecosystem by '
        'digitizing the invitation process. Traditional paper invitations are costly, time-consuming, '
        'and environmentally unfriendly. VibeLink Ghana provides an eco-friendly, cost-effective, '
        'and efficient alternative that allows event organizers to create beautiful digital '
        'invitations, track RSVPs, and manage guest lists seamlessly.'
    )

    doc.add_heading('Target Market', level=2)
    market_segments = [
        ('Individual Event Planners', 'Couples planning weddings, families organizing funerals, '
         'parents hosting naming ceremonies and graduations'),
        ('Corporate Clients', 'Companies organizing product launches, conferences, team building '
         'events, and corporate celebrations'),
        ('Event Planning Agencies', 'Professional event planners managing multiple clients and '
         'events simultaneously'),
        ('Educational Institutions', 'Schools and universities for graduation ceremonies, '
         'award nights, and institutional events')
    ]

    for segment, description in market_segments:
        doc.add_paragraph().add_run(f'{segment}: ').bold = True
        doc.paragraphs[-1].add_run(description)
        doc.paragraphs[-1].paragraph_format.line_spacing = 1.15

    doc.add_heading('Unique Value Proposition', level=2)
    uvp_points = [
        'Culturally Relevant: Designed specifically for Ghanaian events with appropriate templates '
        'and cultural considerations',
        'Affordable Pricing: Competitive pricing structure accessible to various market segments',
        'Easy to Use: Intuitive interface requiring no technical expertise',
        'Instant Delivery: Digital invitations delivered immediately via email and WhatsApp',
        'Track Engagement: Real-time tracking of invitation views and RSVP responses',
        'Professional Templates: Beautifully designed templates for all event types',
        'Secure Payments: Integrated with Paystack for safe and reliable payment processing',
        'Customer Support: AI-powered chatbot for 24/7 assistance'
    ]

    for point in uvp_points:
        p = doc.add_paragraph(point, style='List Bullet')
        p.paragraph_format.line_spacing = 1.15

    doc.add_heading('Website Access', level=2)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'

    table.cell(0, 0).text = 'Website URL'
    table.cell(0, 1).text = 'https://vibelinkgh.com/'
    table.cell(1, 0).text = 'GitHub Repository'
    table.cell(1, 1).text = 'https://github.com/xboggg/vibe-link-ghana-3e5cfc20'
    table.cell(2, 0).text = 'Production Server'
    table.cell(2, 1).text = '/var/www/vibelinkgh.com/'

    doc.add_page_break()

def add_technical_architecture(doc):
    """Add technical architecture section"""
    doc.add_heading('TECHNICAL ARCHITECTURE', level=1)

    doc.add_heading('System Architecture Overview', level=2)
    doc.add_paragraph(
        'VibeLink Ghana follows a modern, cloud-based architecture pattern with clear separation '
        'of concerns between frontend presentation, backend business logic, and data persistence.'
    )

    # Architecture diagram description
    doc.add_paragraph().add_run('[DIAGRAM: System Architecture]').bold = True
    diagram_desc = doc.add_paragraph()
    diagram_desc.add_run('Create a diagram showing:\n')
    diagram_points = [
        'Client Layer: Web Browser (React SPA)',
        'CDN Layer: Static asset delivery',
        'Application Layer: Vite-built React application',
        'API Layer: Supabase Edge Functions',
        'Database Layer: PostgreSQL (Supabase)',
        'External Integrations: Paystack, WhatsApp, Gemini AI, Email Services',
        'Storage: Supabase Storage for images and documents'
    ]
    for point in diagram_points:
        p = doc.add_paragraph(point, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_heading('Technology Stack Breakdown', level=2)

    # Frontend Stack Table
    doc.add_heading('Frontend Technologies', level=3)
    frontend_table = doc.add_table(rows=9, cols=3)
    frontend_table.style = 'Light List Accent 1'

    frontend_data = [
        ('Technology', 'Version', 'Purpose'),
        ('React', '18.3.1', 'UI framework for building component-based interfaces'),
        ('TypeScript', '5.8.3', 'Type safety and enhanced developer experience'),
        ('Vite', '5.4.19', 'Fast build tool and development server'),
        ('Tailwind CSS', '3.4.17', 'Utility-first CSS framework'),
        ('shadcn-ui', 'Latest', 'Accessible UI component library'),
        ('React Router', '6.30.1', 'Client-side routing'),
        ('TanStack Query', '5.83.0', 'Data fetching and state management'),
        ('Framer Motion', '12.23.26', 'Animation library')
    ]

    for i, (tech, version, purpose) in enumerate(frontend_data):
        frontend_table.cell(i, 0).text = tech
        frontend_table.cell(i, 1).text = version
        frontend_table.cell(i, 2).text = purpose

    # Backend Stack Table
    doc.add_heading('Backend Technologies', level=3)
    backend_table = doc.add_table(rows=7, cols=2)
    backend_table.style = 'Light List Accent 1'

    backend_data = [
        ('Technology', 'Purpose'),
        ('Supabase', 'Backend-as-a-Service platform'),
        ('PostgreSQL 14.1', 'Relational database'),
        ('Supabase Auth', 'Authentication and authorization'),
        ('Supabase Edge Functions', 'Serverless functions'),
        ('Supabase Storage', 'File storage'),
        ('Row Level Security (RLS)', 'Database-level security')
    ]

    for i, (tech, purpose) in enumerate(backend_data):
        backend_table.cell(i, 0).text = tech
        backend_table.cell(i, 1).text = purpose

    # Integration Stack
    doc.add_heading('Third-Party Integrations', level=3)
    integration_table = doc.add_table(rows=5, cols=2)
    integration_table.style = 'Light List Accent 1'

    integration_data = [
        ('Service', 'Purpose'),
        ('Paystack', 'Payment processing for Ghana'),
        ('Google Gemini AI', 'AI chatbot for customer support'),
        ('Email Service', 'Transactional emails and notifications'),
        ('WhatsApp Business API', 'Direct WhatsApp invitation delivery')
    ]

    for i, (service, purpose) in enumerate(integration_data):
        integration_table.cell(i, 0).text = service
        integration_table.cell(i, 1).text = purpose

    doc.add_heading('Frontend Architecture', level=2)
    doc.add_paragraph(
        'The frontend follows a component-based architecture with clear separation of concerns:'
    )

    architecture_components = [
        ('Pages (src/pages/)', '32 page components representing different routes'),
        ('Components (src/components/)', 'Reusable UI components organized by feature'),
        ('UI Components (src/components/ui/)', 'Base shadcn-ui components'),
        ('Hooks (src/hooks/)', 'Custom React hooks for shared logic'),
        ('Integrations (src/integrations/)', 'Supabase client and type definitions'),
        ('Data (src/data/)', 'Static data and configuration'),
        ('Assets (src/assets/)', 'Images, icons, and static files')
    ]

    for component, description in architecture_components:
        p = doc.add_paragraph()
        p.add_run(f'{component}: ').bold = True
        p.add_run(description)

    doc.add_heading('Backend Architecture (Supabase)', level=2)
    doc.add_paragraph(
        'Supabase provides a complete backend infrastructure with the following components:'
    )

    backend_components = [
        'PostgreSQL Database with 46 migration files for schema versioning',
        'Automatic REST API generation from database schema',
        'Real-time subscriptions for live data updates',
        'Row Level Security (RLS) policies for fine-grained access control',
        'Edge Functions for custom business logic',
        'Storage buckets for file management',
        'Built-in authentication with email, social providers'
    ]

    for component in backend_components:
        doc.add_paragraph(component, style='List Bullet')

    doc.add_heading('Database Structure Overview', level=2)
    doc.add_paragraph(
        'The database consists of multiple interconnected tables (see Database Schema section '
        'for details). Key tables include:'
    )

    key_tables = [
        'orders: Core order management',
        'customers: Customer information',
        'order_items: Individual items within orders',
        'invitations: Digital invitation data',
        'payments: Payment transaction records',
        'referrals: Referral program tracking',
        'blog_posts: Content management for blog',
        'portfolio_items: Portfolio showcase',
        'surveys: Customer feedback collection',
        'abandoned_carts: Cart abandonment tracking',
        'ai_generated_content: AI-generated content management',
        'analytics_events: User interaction tracking'
    ]

    for table in key_tables:
        doc.add_paragraph(table, style='List Bullet')

    doc.add_heading('Deployment Architecture', level=2)
    doc.add_paragraph(
        'The deployment follows a modern CI/CD pattern with three distinct environments:'
    )

    # Deployment table
    deploy_table = doc.add_table(rows=4, cols=3)
    deploy_table.style = 'Medium Grid 1 Accent 1'

    deploy_data = [
        ('Environment', 'Location', 'Purpose'),
        ('Development', 'C:\\Users\\CyberAware\\OneDrive\\...\\vibelink\\app\\',
         'Local development and testing'),
        ('Version Control', 'https://github.com/xboggg/vibe-link-ghana-3e5cfc20',
         'Central repository and collaboration'),
        ('Production', '/var/www/vibelinkgh.com/', 'Live production server')
    ]

    for i, (env, loc, purpose) in enumerate(deploy_data):
        deploy_table.cell(i, 0).text = env
        deploy_table.cell(i, 1).text = loc
        deploy_table.cell(i, 2).text = purpose

    doc.add_page_break()

def add_development_workflow(doc):
    """Add development workflow section"""
    doc.add_heading('DEVELOPMENT WORKFLOW', level=1)

    doc.add_heading('Three-Location Setup', level=2)
    doc.add_paragraph(
        'The development workflow is organized across three strategic locations, each serving '
        'a specific purpose in the software development lifecycle:'
    )

    doc.add_heading('1. Local Development (PC)', level=3)
    pc_details = [
        'Location: C:\\Users\\CyberAware\\OneDrive - Government of Ghana - CAGD\\ZeroTrust\\Visual Studio Code Workspace\\vibelink\\app\\',
        'Purpose: Active development, testing, and debugging',
        'Tools: Visual Studio Code, Node.js, npm/bun, Git',
        'Server: Vite dev server on localhost:8080',
        'Hot Module Replacement (HMR) for instant feedback',
        'Access to .env file for environment configuration'
    ]

    for detail in pc_details:
        doc.add_paragraph(detail, style='List Bullet')

    doc.add_heading('2. Version Control (GitHub)', level=3)
    github_details = [
        'Repository: https://github.com/xboggg/vibe-link-ghana-3e5cfc20',
        'Purpose: Code versioning, collaboration, and backup',
        'Branch Strategy: Main branch for production-ready code',
        'Commit History: Complete audit trail of changes',
        'Collaboration: Code reviews and team coordination',
        'CI/CD Integration: Automated testing and deployment triggers'
    ]

    for detail in github_details:
        doc.add_paragraph(detail, style='List Bullet')

    doc.add_heading('3. Production Server', level=3)
    server_details = [
        'Location: /var/www/vibelinkgh.com/',
        'Purpose: Live production environment',
        'Web Server: Nginx/Apache serving static files',
        'Domain: https://vibelinkgh.com/',
        'SSL Certificate: HTTPS encryption enabled',
        'Optimization: Minified and bundled production build'
    ]

    for detail in server_details:
        doc.add_paragraph(detail, style='List Bullet')

    doc.add_heading('Git Workflow Diagram', level=2)
    doc.add_paragraph().add_run('[DIAGRAM: Git Workflow]').bold = True
    workflow_desc = doc.add_paragraph()
    workflow_desc.add_run('Create a flowchart showing:\n')
    workflow_steps = [
        '1. Developer writes code on Local PC',
        '2. Git add and commit changes locally',
        '3. Git push to GitHub repository',
        '4. GitHub receives and stores code',
        '5. GitHub Actions or deployment script triggers',
        '6. Build process: npm run build creates optimized dist/ folder',
        '7. Deployment script transfers dist/ contents to server',
        '8. Server updates /var/www/vibelinkgh.com/',
        '9. Website live with latest changes'
    ]

    for step in workflow_steps:
        p = doc.add_paragraph(step, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_heading('Build and Deployment Process', level=2)

    doc.add_heading('Development Build', level=3)
    doc.add_paragraph('Command: npm run dev or bun run dev')
    dev_process = [
        'Vite starts development server on port 8080',
        'Hot Module Replacement (HMR) enabled',
        'Source maps for debugging',
        'Fast refresh for instant updates',
        'Environment: Development mode'
    ]
    for item in dev_process:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Production Build', level=3)
    doc.add_paragraph('Command: npm run build or bun run build')
    prod_process = [
        'TypeScript compilation and type checking',
        'React component tree shaking',
        'CSS purging and minification',
        'JavaScript minification and uglification',
        'Asset optimization (images, fonts)',
        'Bundle splitting for optimal loading',
        'Output: dist/ directory with production-ready files',
        'Build time: ~30-60 seconds'
    ]
    for item in prod_process:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Deployment Steps', level=3)
    deploy_steps = [
        'Build the production bundle locally or in CI/CD',
        'Test the build with: npm run preview',
        'Transfer dist/ contents to /var/www/vibelinkgh.com/',
        'Update file permissions on server',
        'Clear server cache if applicable',
        'Verify deployment by visiting vibelinkgh.com',
        'Monitor error logs for any issues'
    ]

    for i, step in enumerate(deploy_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('Environment Setup', level=2)

    doc.add_heading('Prerequisites', level=3)
    prereqs = [
        'Node.js (v18 or higher)',
        'npm or bun package manager',
        'Git for version control',
        'Visual Studio Code (recommended IDE)',
        'Supabase account and project',
        'Paystack merchant account'
    ]
    for prereq in prereqs:
        doc.add_paragraph(prereq, style='List Bullet')

    doc.add_heading('Initial Setup', level=3)
    setup_code = '''
1. Clone repository:
   git clone https://github.com/xboggg/vibe-link-ghana-3e5cfc20

2. Navigate to project:
   cd vibe-link-ghana-3e5cfc20/app

3. Install dependencies:
   npm install
   (or: bun install)

4. Configure environment:
   Create .env file with:
   VITE_SUPABASE_URL=your_supabase_url
   VITE_SUPABASE_PUBLISHABLE_KEY=your_supabase_key
   VITE_SUPABASE_PROJECT_ID=your_project_id

5. Start development server:
   npm run dev

6. Open browser to:
   http://localhost:8080
'''

    p = doc.add_paragraph(setup_code)
    p.style = 'No Spacing'

    doc.add_page_break()

def add_features_specifications(doc):
    """Add features and specifications section"""
    doc.add_heading('FEATURES & SPECIFICATIONS', level=1)

    doc.add_heading('Customer-Facing Features', level=2)

    doc.add_heading('1. Digital Invitation Creation', level=3)
    doc.add_paragraph(
        'Customers can create beautiful, customized digital invitations through an intuitive '
        'form-based interface:'
    )
    invitation_features = [
        'Event type selection (Wedding, Funeral, Naming Ceremony, Graduation, Corporate)',
        'Template gallery with culturally appropriate designs',
        'Custom text and messaging',
        'Event details: date, time, venue with map integration',
        'Guest list management',
        'RSVP tracking',
        'Photo upload for personalization',
        'Color scheme customization',
        'Font selection',
        'Preview before finalizing'
    ]
    for feature in invitation_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('2. Order Management', level=3)
    order_features = [
        'Order placement through "Get Started" page',
        'Package selection (Basic, Standard, Premium)',
        'Add-on services selection',
        'Real-time price calculation',
        'Order summary review',
        'Secure payment processing',
        'Order confirmation email',
        'Order tracking with unique ID',
        'Order history in customer portal',
        'Invoice generation and download'
    ]
    for feature in order_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('3. Customer Portal', level=3)
    doc.add_paragraph('Authenticated customers access a personalized dashboard featuring:')
    portal_features = [
        'Overview of all orders',
        'Active orders with status tracking',
        'Completed orders archive',
        'Order details view',
        'Invoice download',
        'Payment history',
        'Profile management',
        'Saved templates',
        'Referral code and earnings',
        'Support ticket system'
    ]
    for feature in portal_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('4. Order Tracking', level=3)
    tracking_features = [
        'Track order by Order ID or Email',
        'Real-time status updates',
        'Status stages: Pending, Processing, Design Review, Completed',
        'Estimated completion time',
        'Email notifications on status changes',
        'WhatsApp notifications (optional)',
        'Design preview when ready',
        'Revision request capability'
    ]
    for feature in tracking_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('5. Referral Program', level=3)
    doc.add_paragraph(
        'Customers can earn rewards by referring others to VibeLink Ghana:'
    )
    referral_features = [
        'Unique referral code generation',
        'Shareable referral link',
        'Social media sharing buttons',
        'Track referrals and conversions',
        'Earnings dashboard',
        'Automatic commission calculation',
        'Withdrawal request system',
        'Referral leaderboard',
        'Bonus incentives for milestones'
    ]
    for feature in referral_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('Admin Features', level=2)

    doc.add_heading('1. Admin Dashboard', level=3)
    doc.add_paragraph('Comprehensive admin panel with real-time analytics and management:')
    admin_dashboard = [
        'Overview statistics (orders, revenue, customers)',
        'Real-time order notifications',
        'Quick action buttons',
        'Recent orders list',
        'Revenue charts and graphs',
        'Customer growth metrics',
        'Popular packages analysis',
        'Peak usage times',
        'Conversion rate tracking'
    ]
    for feature in admin_dashboard:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('2. Order Management System', level=3)
    admin_orders = [
        'All orders view with filtering',
        'Order status management',
        'Customer information access',
        'Design upload and management',
        'Revision handling',
        'Communication with customers',
        'Order priority setting',
        'Bulk actions',
        'Export to CSV/Excel',
        'Print packing slips'
    ]
    for feature in admin_orders:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('3. Customer Management', level=3)
    admin_customers = [
        'Customer database with search',
        'Customer profiles and history',
        'Lifetime value calculation',
        'Customer segmentation',
        'Email customer directly',
        'Add notes to customer profiles',
        'Block/unblock customers',
        'Export customer data',
        'Customer activity logs'
    ]
    for feature in admin_customers:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('4. Content Management', level=3)
    admin_content = [
        'Blog post creation and editing (with TipTap rich text editor)',
        'Portfolio item management',
        'Image upload and gallery',
        'SEO meta data editing',
        'Publish/draft status',
        'Scheduling capabilities',
        'Categories and tags',
        'Featured posts',
        'Analytics per post'
    ]
    for feature in admin_content:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('5. Analytics & Reports', level=3)
    admin_analytics = [
        'Revenue reports (daily, weekly, monthly)',
        'Order statistics',
        'Customer acquisition reports',
        'Referral program performance',
        'Popular templates/packages',
        'Conversion funnel analysis',
        'Abandoned cart reports',
        'Traffic source analytics',
        'User behavior tracking',
        'Export reports to PDF/Excel'
    ]
    for feature in admin_analytics:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('Payment Processing', level=2)
    doc.add_paragraph(
        'Integration with Paystack provides secure and reliable payment processing:'
    )
    payment_features = [
        'Paystack payment gateway integration',
        'Support for mobile money (MTN, Vodafone, AirtelTigo)',
        'Card payments (Visa, Mastercard)',
        'Bank transfer option',
        'Secure payment page (PCI DSS compliant)',
        'Automatic payment verification',
        'Payment confirmation emails',
        'Transaction history',
        'Refund processing',
        'Payment webhooks for real-time updates',
        'Multi-currency support (GHS primary)',
        'Test mode for development'
    ]
    for feature in payment_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('AI Chatbot Integration', level=2)
    doc.add_paragraph(
        'Google Gemini AI powers an intelligent chatbot for customer support:'
    )
    chatbot_features = [
        'Natural language understanding',
        '24/7 availability',
        'Instant responses to common questions',
        'Multi-language support',
        'Contextual conversations',
        'FAQs and knowledge base integration',
        'Escalation to human support',
        'Order status inquiries',
        'Package recommendations',
        'Troubleshooting assistance',
        'Chat history',
        'Feedback collection'
    ]
    for feature in chatbot_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('SEO Features', level=2)
    seo_features = [
        'React Helmet for dynamic meta tags',
        'SEO-friendly URLs',
        'Open Graph tags for social sharing',
        'Twitter Card integration',
        'Sitemap generation',
        'Robots.txt configuration',
        'Canonical URLs',
        'Structured data (Schema.org)',
        'Image optimization with alt tags',
        'Fast page load times',
        'Mobile responsiveness',
        'HTTPS security',
        'Internal linking structure'
    ]
    for feature in seo_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('Integration Points', level=2)

    # Integration table
    integration_table = doc.add_table(rows=6, cols=3)
    integration_table.style = 'Light Grid Accent 1'

    integration_data = [
        ('Integration', 'Technology', 'Purpose'),
        ('Payment Gateway', 'Paystack API', 'Process payments securely'),
        ('AI Chatbot', 'Google Gemini AI', 'Intelligent customer support'),
        ('Email Service', 'Supabase Email/SMTP', 'Transactional emails'),
        ('WhatsApp', 'WhatsApp Business API', 'Direct message delivery'),
        ('Analytics', 'Custom + Google Analytics', 'User behavior tracking')
    ]

    for i, (integration, tech, purpose) in enumerate(integration_data):
        integration_table.cell(i, 0).text = integration
        integration_table.cell(i, 1).text = tech
        integration_table.cell(i, 2).text = purpose

    doc.add_page_break()

def add_user_interface(doc):
    """Add user interface section"""
    doc.add_heading('USER INTERFACE', level=1)

    doc.add_heading('Page Descriptions', level=2)

    # Public Pages
    doc.add_heading('Public Pages', level=3)
    public_pages = [
        ('Index (Home)', 'Landing page with hero section, features showcase, testimonials, '
         'pricing preview, call-to-action buttons, and recent portfolio items'),
        ('About', 'Company story, mission, vision, team members, and values'),
        ('Services', 'Detailed description of invitation services, event types covered, '
         'customization options, and delivery methods'),
        ('Pricing', 'Package comparison (Basic, Standard, Premium), pricing table, '
         'add-on services, and clear call-to-action'),
        ('Portfolio', 'Gallery of past invitation designs, filterable by event type, '
         'searchable, with pagination'),
        ('Portfolio Detail', 'Full view of individual portfolio items with multiple images, '
         'description, and related items'),
        ('How It Works', 'Step-by-step explanation of the process from order to delivery, '
         'with visual timeline'),
        ('Blog', 'Blog listing page with featured posts, categories, search, and pagination'),
        ('Blog Detail', 'Individual blog post with rich content, sharing buttons, comments, '
         'and related posts'),
        ('Contact', 'Contact form, office location map, phone numbers, email, social media links'),
        ('Get Started', 'Order form with package selection, event details input, '
         'customization options, and checkout'),
        ('Track Order', 'Order tracking by ID or email, status display, timeline view')
    ]

    for page_name, description in public_pages:
        p = doc.add_paragraph()
        p.add_run(f'{page_name}: ').bold = True
        p.add_run(description)
        p.paragraph_format.line_spacing = 1.15

    # Customer Portal Pages
    doc.add_heading('Customer Portal Pages (Authenticated)', level=3)
    customer_pages = [
        ('Customer Portal Dashboard', 'Overview of orders, recent activity, quick stats, '
         'referral earnings, and navigation to other portal sections'),
        ('Order Details', 'Complete order information, status timeline, design preview, '
         'invoice download, revision request'),
        ('Invoice View', 'Professional invoice with company details, order breakdown, '
         'payment information, printable format')
    ]

    for page_name, description in customer_pages:
        p = doc.add_paragraph()
        p.add_run(f'{page_name}: ').bold = True
        p.add_run(description)
        p.paragraph_format.line_spacing = 1.15

    # Admin Pages
    doc.add_heading('Admin Pages (Admin Access Only)', level=3)
    admin_pages = [
        ('Admin Auth', 'Secure admin login with email/password, optional 2FA'),
        ('Admin Dashboard', 'Comprehensive admin overview with analytics, charts, '
         'recent orders, quick actions, and system health')
    ]

    for page_name, description in admin_pages:
        p = doc.add_paragraph()
        p.add_run(f'{page_name}: ').bold = True
        p.add_run(description)
        p.paragraph_format.line_spacing = 1.15

    # Legal Pages
    doc.add_heading('Legal & Policy Pages', level=3)
    legal_pages = [
        ('Privacy Policy', 'Data collection, usage, protection, and user rights'),
        ('Terms of Service', 'User agreements, service terms, limitations, and disclaimers'),
        ('Cookie Policy', 'Cookie usage, types, and user controls'),
        ('Refund Policy', 'Refund conditions, process, and timelines')
    ]

    for page_name, description in legal_pages:
        p = doc.add_paragraph()
        p.add_run(f'{page_name}: ').bold = True
        p.add_run(description)
        p.paragraph_format.line_spacing = 1.15

    # Other Pages
    doc.add_heading('Other Pages', level=3)
    other_pages = [
        ('Survey', 'Customer feedback and satisfaction survey'),
        ('Thank You', 'Order confirmation and thank you message'),
        ('Multi-Language Support', 'Language selection interface'),
        ('Not Found (404)', 'Custom 404 error page with navigation options')
    ]

    for page_name, description in other_pages:
        p = doc.add_paragraph()
        p.add_run(f'{page_name}: ').bold = True
        p.add_run(description)
        p.paragraph_format.line_spacing = 1.15

    doc.add_heading('User Journeys', level=2)

    doc.add_heading('Customer Journey: Ordering an Invitation', level=3)
    customer_journey = [
        '1. Visitor lands on homepage (Index.tsx)',
        '2. Explores services and views portfolio',
        '3. Checks pricing packages',
        '4. Clicks "Get Started" CTA button',
        '5. Fills out order form with event details',
        '6. Selects package and customization options',
        '7. Reviews order summary',
        '8. Proceeds to payment',
        '9. Completes payment via Paystack',
        '10. Receives order confirmation email',
        '11. Gets order tracking ID',
        '12. Can track order status',
        '13. Receives design for review',
        '14. Approves or requests revisions',
        '15. Receives final invitation',
        '16. Downloads invoice from customer portal'
    ]

    for step in customer_journey:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('Admin Journey: Processing an Order', level=3)
    admin_journey = [
        '1. Admin logs in via AdminAuth.tsx',
        '2. Sees new order notification on dashboard',
        '3. Clicks to view order details',
        '4. Reviews customer requirements',
        '5. Changes order status to "Processing"',
        '6. Creates invitation design',
        '7. Uploads design to order',
        '8. Changes status to "Design Review"',
        '9. Customer receives notification',
        '10. Customer reviews and provides feedback',
        '11. Admin makes revisions if needed',
        '12. Customer approves final design',
        '13. Admin marks order as "Completed"',
        '14. System sends final invitation to customer',
        '15. Admin monitors customer satisfaction'
    ]

    for step in admin_journey:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('Navigation Structure', level=2)
    doc.add_paragraph('The site navigation is organized hierarchically:')

    doc.add_paragraph().add_run('Main Navigation:').bold = True
    main_nav = ['Home', 'About', 'Services', 'Portfolio', 'Pricing', 'Blog',
                'How It Works', 'Contact', 'Get Started', 'Track Order']
    for item in main_nav:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph().add_run('Footer Navigation:').bold = True
    footer_nav = ['About Us', 'Services', 'Contact', 'Privacy Policy', 'Terms of Service',
                  'Cookie Policy', 'Refund Policy', 'Blog', 'FAQ']
    for item in footer_nav:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph().add_run('Customer Portal Navigation:').bold = True
    portal_nav = ['Dashboard', 'My Orders', 'Profile', 'Referrals', 'Support', 'Logout']
    for item in portal_nav:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph().add_run('Admin Panel Navigation:').bold = True
    admin_nav = ['Dashboard', 'Orders', 'Customers', 'Content', 'Analytics',
                 'Settings', 'Logout']
    for item in admin_nav:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_heading('Design System', level=2)

    doc.add_heading('Color Palette', level=3)
    colors = [
        ('Navy Blue', '#1A237E', 'Primary brand color, headers, CTAs'),
        ('Gold', '#FFD700', 'Accent color, highlights, premium features'),
        ('Purple', '#7C3AED', 'Secondary accent, interactive elements'),
        ('Mint', '#E8F5E9', 'Success states, positive feedback'),
        ('White', '#FFFFFF', 'Backgrounds, cards'),
        ('Gray Scale', '#F5F5F5 to #212121', 'Text, borders, backgrounds')
    ]

    color_table = doc.add_table(rows=len(colors) + 1, cols=3)
    color_table.style = 'Light Grid Accent 1'
    color_table.cell(0, 0).text = 'Color Name'
    color_table.cell(0, 1).text = 'Hex Code'
    color_table.cell(0, 2).text = 'Usage'

    for i, (name, hex_code, usage) in enumerate(colors, 1):
        color_table.cell(i, 0).text = name
        color_table.cell(i, 1).text = hex_code
        color_table.cell(i, 2).text = usage

    doc.add_heading('Typography', level=3)
    fonts = [
        ('Poppins', 'Primary font for headings and titles', '600-700 weight'),
        ('Inter', 'Body text and UI elements', '400-500 weight'),
        ('Playfair Display', 'Decorative headings and special sections', '700 weight')
    ]

    font_table = doc.add_table(rows=len(fonts) + 1, cols=3)
    font_table.style = 'Light Grid Accent 1'
    font_table.cell(0, 0).text = 'Font Family'
    font_table.cell(0, 1).text = 'Usage'
    font_table.cell(0, 2).text = 'Weights'

    for i, (family, usage, weights) in enumerate(fonts, 1):
        font_table.cell(i, 0).text = family
        font_table.cell(i, 1).text = usage
        font_table.cell(i, 2).text = weights

    doc.add_heading('Spacing System', level=3)
    doc.add_paragraph('Consistent spacing using Tailwind CSS spacing scale:')
    spacing = ['4px (1)', '8px (2)', '12px (3)', '16px (4)', '24px (6)',
               '32px (8)', '48px (12)', '64px (16)', '96px (24)']
    for space in spacing:
        doc.add_paragraph(space, style='List Bullet')

    doc.add_heading('Component Library', level=3)
    doc.add_paragraph('UI components built with shadcn-ui and Radix UI:')
    components = [
        'Buttons (primary, secondary, outline, ghost)',
        'Forms (inputs, selects, textareas, checkboxes, radio buttons)',
        'Cards (content containers with shadows)',
        'Modals and Dialogs',
        'Dropdowns and Popovers',
        'Tables (responsive data tables)',
        'Tabs and Accordions',
        'Toasts and Notifications',
        'Progress Indicators',
        'Avatars and Badges',
        'Navigation (navbar, sidebar, breadcrumbs)',
        'Charts and Graphs (Recharts)'
    ]
    for component in components:
        doc.add_paragraph(component, style='List Bullet')

    doc.add_page_break()

def add_database_schema(doc):
    """Add database schema section"""
    doc.add_heading('DATABASE SCHEMA', level=1)

    doc.add_heading('Database Overview', level=2)
    doc.add_paragraph(
        'VibeLink Ghana uses PostgreSQL 14.1 via Supabase with 46 migration files tracking '
        'schema evolution. The database implements Row Level Security (RLS) for fine-grained '
        'access control and uses foreign key constraints to maintain referential integrity.'
    )

    doc.add_heading('Core Tables', level=2)

    # Orders Table
    doc.add_heading('orders', level=3)
    doc.add_paragraph('Core table for managing customer orders.')
    orders_fields = [
        ('id', 'UUID', 'Primary key, auto-generated'),
        ('customer_id', 'UUID', 'Foreign key to customers table'),
        ('order_number', 'VARCHAR', 'Unique human-readable order ID'),
        ('event_type', 'VARCHAR', 'Wedding, Funeral, Naming, Graduation, Corporate'),
        ('package_name', 'VARCHAR', 'Basic, Standard, Premium'),
        ('status', 'VARCHAR', 'Pending, Processing, Design Review, Completed, Cancelled'),
        ('total_price', 'DECIMAL', 'Total order amount in GHS'),
        ('payment_status', 'VARCHAR', 'Pending, Paid, Refunded'),
        ('created_at', 'TIMESTAMP', 'Order creation timestamp'),
        ('updated_at', 'TIMESTAMP', 'Last update timestamp'),
        ('completed_at', 'TIMESTAMP', 'Order completion timestamp'),
        ('notes', 'TEXT', 'Admin notes and special instructions')
    ]

    orders_table = doc.add_table(rows=len(orders_fields) + 1, cols=3)
    orders_table.style = 'Medium Grid 1 Accent 1'
    orders_table.cell(0, 0).text = 'Field'
    orders_table.cell(0, 1).text = 'Type'
    orders_table.cell(0, 2).text = 'Description'

    for i, (field, ftype, desc) in enumerate(orders_fields, 1):
        orders_table.cell(i, 0).text = field
        orders_table.cell(i, 1).text = ftype
        orders_table.cell(i, 2).text = desc

    # Customers Table
    doc.add_heading('customers', level=3)
    doc.add_paragraph('Stores customer account information.')
    customers_fields = [
        ('id', 'UUID', 'Primary key, matches Supabase auth.users.id'),
        ('email', 'VARCHAR', 'Customer email (unique)'),
        ('full_name', 'VARCHAR', 'Customer full name'),
        ('phone', 'VARCHAR', 'Phone number'),
        ('created_at', 'TIMESTAMP', 'Account creation date'),
        ('updated_at', 'TIMESTAMP', 'Last profile update'),
        ('total_orders', 'INTEGER', 'Count of orders placed'),
        ('lifetime_value', 'DECIMAL', 'Total revenue from customer'),
        ('referral_code', 'VARCHAR', 'Unique referral code'),
        ('referred_by', 'UUID', 'Foreign key to customers (referrer)')
    ]

    customers_table = doc.add_table(rows=len(customers_fields) + 1, cols=3)
    customers_table.style = 'Medium Grid 1 Accent 1'
    customers_table.cell(0, 0).text = 'Field'
    customers_table.cell(0, 1).text = 'Type'
    customers_table.cell(0, 2).text = 'Description'

    for i, (field, ftype, desc) in enumerate(customers_fields, 1):
        customers_table.cell(i, 0).text = field
        customers_table.cell(i, 1).text = ftype
        customers_table.cell(i, 2).text = desc

    # Payments Table
    doc.add_heading('payments', level=3)
    doc.add_paragraph('Tracks all payment transactions.')
    payments_fields = [
        ('id', 'UUID', 'Primary key'),
        ('order_id', 'UUID', 'Foreign key to orders'),
        ('amount', 'DECIMAL', 'Payment amount'),
        ('currency', 'VARCHAR', 'GHS, USD, etc.'),
        ('status', 'VARCHAR', 'Success, Pending, Failed'),
        ('payment_method', 'VARCHAR', 'Card, Mobile Money, Bank Transfer'),
        ('provider_reference', 'VARCHAR', 'Paystack transaction reference'),
        ('created_at', 'TIMESTAMP', 'Payment initiation time'),
        ('completed_at', 'TIMESTAMP', 'Payment completion time')
    ]

    payments_table = doc.add_table(rows=len(payments_fields) + 1, cols=3)
    payments_table.style = 'Medium Grid 1 Accent 1'
    payments_table.cell(0, 0).text = 'Field'
    payments_table.cell(0, 1).text = 'Type'
    payments_table.cell(0, 2).text = 'Description'

    for i, (field, ftype, desc) in enumerate(payments_fields, 1):
        payments_table.cell(i, 0).text = field
        payments_table.cell(i, 1).text = ftype
        payments_table.cell(i, 2).text = desc

    doc.add_heading('Supporting Tables', level=2)

    supporting_tables = [
        ('order_items', 'Individual line items within an order (invitation designs, add-ons)'),
        ('invitations', 'Stores invitation design data, templates, and customizations'),
        ('referrals', 'Tracks referral relationships and commission earnings'),
        ('blog_posts', 'Blog content with title, content, author, publish date, SEO data'),
        ('portfolio_items', 'Portfolio showcase items with images and descriptions'),
        ('surveys', 'Customer feedback surveys and responses'),
        ('abandoned_carts', 'Tracks incomplete orders for recovery campaigns'),
        ('ai_generated_content', 'AI-generated content for approval workflow'),
        ('analytics_events', 'User interaction tracking for behavioral analysis'),
        ('email_templates', 'Transactional email templates'),
        ('notifications', 'System notifications for customers and admins'),
        ('rate_limits', 'API rate limiting data'),
        ('admin_users', 'Admin account information and permissions')
    ]

    for table_name, description in supporting_tables:
        p = doc.add_paragraph()
        p.add_run(f'{table_name}: ').bold = True
        p.add_run(description)

    doc.add_heading('Key Relationships', level=2)
    doc.add_paragraph('Database relationships enforce data integrity:')

    relationships = [
        'customers 1:N orders (one customer can have many orders)',
        'orders 1:N order_items (one order can have multiple items)',
        'orders 1:N payments (one order can have multiple payment attempts)',
        'customers 1:1 referral_code (unique referral code per customer)',
        'customers 1:N referrals (one customer can refer many others)',
        'orders 1:1 invitations (each order has one invitation design)',
        'customers 1:N notifications (one customer receives many notifications)',
        'customers 1:N abandoned_carts (tracking multiple cart abandonment events)'
    ]

    for relationship in relationships:
        doc.add_paragraph(relationship, style='List Bullet')

    doc.add_heading('Migration Strategy', level=2)
    doc.add_paragraph(
        'Database schema changes are managed through Supabase migrations with 46 migration '
        'files tracking the evolution of the database schema.'
    )

    doc.add_heading('Migration Process', level=3)
    migration_steps = [
        'Create new migration file: supabase migration new migration_name',
        'Write SQL DDL statements in migration file',
        'Test migration in local development environment',
        'Review migration for potential data loss or performance issues',
        'Apply migration to staging environment',
        'Verify migration success and data integrity',
        'Apply migration to production environment',
        'Monitor for errors and performance degradation',
        'Document migration in changelog'
    ]

    for i, step in enumerate(migration_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('Example Migration Files', level=3)
    example_migrations = [
        '20241230_order_priority.sql - Added priority field to orders table',
        '20241230_rate_limit_table.sql - Created rate limiting table',
        '20251225000923_*.sql - Various schema enhancements'
    ]

    for migration in example_migrations:
        doc.add_paragraph(migration, style='List Bullet')

    doc.add_heading('Row Level Security (RLS)', level=3)
    doc.add_paragraph(
        'RLS policies ensure users can only access data they are authorized to see:'
    )

    rls_policies = [
        'Customers can only view their own orders and profile',
        'Admins have full access to all data',
        'Public users can view portfolio and blog posts',
        'Unauthenticated users cannot access customer or order data',
        'Insert policies validate data before allowing creation',
        'Update policies ensure users can only modify their own data',
        'Delete policies restrict data deletion to authorized users'
    ]

    for policy in rls_policies:
        doc.add_paragraph(policy, style='List Bullet')

    doc.add_page_break()

def add_api_integrations(doc):
    """Add API and integrations section"""
    doc.add_heading('API & INTEGRATIONS', level=1)

    doc.add_heading('Supabase API Usage', level=2)
    doc.add_paragraph(
        'Supabase provides auto-generated RESTful API and real-time subscriptions for all '
        'database tables. The frontend communicates with Supabase using the official '
        'JavaScript client library.'
    )

    doc.add_heading('Authentication API', level=3)
    auth_endpoints = [
        ('signUp()', 'Create new user account with email/password'),
        ('signIn()', 'Authenticate user and receive session token'),
        ('signOut()', 'End user session'),
        ('resetPassword()', 'Send password reset email'),
        ('updateUser()', 'Update user profile information'),
        ('getSession()', 'Retrieve current session information'),
        ('onAuthStateChange()', 'Subscribe to authentication state changes')
    ]

    for method, description in auth_endpoints:
        p = doc.add_paragraph()
        p.add_run(f'{method}: ').bold = True
        p.add_run(description)

    doc.add_heading('Database API', level=3)
    doc.add_paragraph('CRUD operations via Supabase client:')

    crud_example = '''
// Query orders
const { data, error } = await supabase
  .from('orders')
  .select('*')
  .eq('customer_id', userId)
  .order('created_at', { ascending: false })

// Insert new order
const { data, error } = await supabase
  .from('orders')
  .insert({
    customer_id: userId,
    event_type: 'wedding',
    package_name: 'premium',
    total_price: 500
  })

// Update order status
const { data, error } = await supabase
  .from('orders')
  .update({ status: 'completed' })
  .eq('id', orderId)

// Delete (soft delete preferred)
const { data, error } = await supabase
  .from('orders')
  .update({ deleted_at: new Date() })
  .eq('id', orderId)
'''

    p = doc.add_paragraph(crud_example)
    p.style = 'No Spacing'

    doc.add_heading('Real-time Subscriptions', level=3)
    doc.add_paragraph('Subscribe to database changes for real-time updates:')

    realtime_example = '''
// Subscribe to new orders
const subscription = supabase
  .channel('orders')
  .on('postgres_changes',
    {
      event: 'INSERT',
      schema: 'public',
      table: 'orders'
    },
    (payload) => {
      console.log('New order:', payload.new)
      // Update UI with new order
    }
  )
  .subscribe()
'''

    p = doc.add_paragraph(realtime_example)
    p.style = 'No Spacing'

    doc.add_heading('Storage API', level=3)
    storage_operations = [
        ('upload()', 'Upload files to Supabase Storage buckets'),
        ('download()', 'Download files from storage'),
        ('createSignedUrl()', 'Generate temporary URL for private files'),
        ('list()', 'List files in a bucket'),
        ('remove()', 'Delete files from storage'),
        ('getPublicUrl()', 'Get public URL for public bucket files')
    ]

    for method, description in storage_operations:
        p = doc.add_paragraph()
        p.add_run(f'{method}: ').bold = True
        p.add_run(description)

    doc.add_heading('Paystack Payment Integration', level=2)
    doc.add_paragraph(
        'Paystack is integrated for payment processing with support for multiple payment '
        'methods popular in Ghana.'
    )

    doc.add_heading('Payment Flow', level=3)
    payment_flow = [
        '1. Customer completes order form',
        '2. Frontend initiates payment with Paystack',
        '3. Paystack popup opens with payment options',
        '4. Customer selects payment method (card/mobile money)',
        '5. Customer completes payment on Paystack',
        '6. Paystack sends webhook to backend',
        '7. Backend verifies payment with Paystack API',
        '8. Order status updated to "Paid"',
        '9. Customer receives confirmation email',
        '10. Admin receives new order notification'
    ]

    for step in payment_flow:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('Paystack API Endpoints', level=3)
    paystack_endpoints = [
        ('Initialize Transaction', 'POST /transaction/initialize - Create payment session'),
        ('Verify Transaction', 'GET /transaction/verify/:reference - Confirm payment'),
        ('List Transactions', 'GET /transaction - Retrieve transaction history'),
        ('Refund Transaction', 'POST /refund - Process refund'),
        ('Mobile Money', 'POST /charge - Charge mobile money accounts')
    ]

    for endpoint, description in paystack_endpoints:
        p = doc.add_paragraph()
        p.add_run(f'{endpoint}: ').bold = True
        p.add_run(description)

    doc.add_heading('Payment Methods Supported', level=3)
    payment_methods = [
        'Visa/Mastercard credit and debit cards',
        'MTN Mobile Money',
        'Vodafone Cash',
        'AirtelTigo Money',
        'Bank transfer',
        'USSD'
    ]

    for method in payment_methods:
        doc.add_paragraph(method, style='List Bullet')

    doc.add_heading('WhatsApp Integration', level=2)
    doc.add_paragraph(
        'WhatsApp Business API integration enables direct delivery of invitations and '
        'notifications to customers.'
    )

    doc.add_heading('WhatsApp Features', level=3)
    whatsapp_features = [
        'Send invitation links directly to customer WhatsApp',
        'Order status notifications',
        'Payment confirmations',
        'Design ready notifications',
        'Support message delivery',
        'Bulk WhatsApp messaging for campaigns',
        'Message templates for compliance',
        'Delivery status tracking'
    ]

    for feature in whatsapp_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('Implementation', level=3)
    doc.add_paragraph(
        'WhatsApp integration uses the WhatsApp Business API with message templates '
        'approved by Meta. Messages are sent via Supabase Edge Functions to maintain '
        'security and comply with WhatsApp policies.'
    )

    doc.add_heading('Google Gemini AI Integration', level=2)
    doc.add_paragraph(
        'Google Gemini AI powers the intelligent chatbot providing 24/7 customer support.'
    )

    doc.add_heading('Chatbot Capabilities', level=3)
    chatbot_capabilities = [
        'Answer frequently asked questions',
        'Explain pricing and packages',
        'Guide users through ordering process',
        'Check order status',
        'Provide design recommendations',
        'Troubleshoot common issues',
        'Escalate to human support when needed',
        'Multi-turn conversations with context',
        'Natural language understanding',
        'Response in English and local languages'
    ]

    for capability in chatbot_capabilities:
        doc.add_paragraph(capability, style='List Bullet')

    doc.add_heading('Implementation', level=3)
    doc.add_paragraph(
        'The chatbot is implemented as a floating widget (ChatWidget.tsx) that appears on '
        'all pages. It communicates with Google Gemini AI via Supabase Edge Functions '
        'to protect API keys.'
    )

    gemini_flow = [
        '1. User types message in chat widget',
        '2. Frontend sends message to Supabase Edge Function',
        '3. Edge Function calls Gemini AI API with message and context',
        '4. Gemini AI processes message and generates response',
        '5. Response sent back to Edge Function',
        '6. Edge Function returns response to frontend',
        '7. Chat widget displays response to user',
        '8. Conversation history maintained for context'
    ]

    for step in gemini_flow:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('Email Services', level=2)
    doc.add_paragraph(
        'Email notifications are sent for various events throughout the customer journey.'
    )

    doc.add_heading('Email Types', level=3)
    email_types = [
        ('Welcome Email', 'Sent when customer creates account'),
        ('Order Confirmation', 'Sent after successful order placement'),
        ('Payment Confirmation', 'Sent after successful payment'),
        ('Order Status Update', 'Sent when order status changes'),
        ('Design Ready', 'Sent when design is ready for review'),
        ('Order Completed', 'Sent when order is finalized'),
        ('Abandoned Cart', 'Reminder for incomplete orders'),
        ('Password Reset', 'Password reset link'),
        ('Invoice', 'Order invoice attached'),
        ('Referral Notifications', 'Referral commission updates')
    ]

    for email_type, description in email_types:
        p = doc.add_paragraph()
        p.add_run(f'{email_type}: ').bold = True
        p.add_run(description)

    doc.add_heading('Email Templates', level=3)
    doc.add_paragraph(
        'Email templates are stored in the database and support dynamic content insertion. '
        'Templates are designed to be mobile-responsive with consistent branding.'
    )

    doc.add_heading('Sending Method', level=3)
    doc.add_paragraph(
        'Emails are sent via Supabase Edge Functions using SMTP or a transactional email '
        'service like SendGrid or Mailgun.'
    )

    doc.add_page_break()

def add_deployment_infrastructure(doc):
    """Add deployment and infrastructure section"""
    doc.add_heading('DEPLOYMENT & INFRASTRUCTURE', level=1)

    doc.add_heading('Server Setup', level=2)
    doc.add_paragraph(
        'The production server hosts the VibeLink Ghana application with optimized '
        'configuration for performance and security.'
    )

    doc.add_heading('Server Specifications', level=3)
    server_specs = [
        'Operating System: Linux (Ubuntu/Debian recommended)',
        'Web Server: Nginx or Apache',
        'Node.js: Not required (static site)',
        'SSL/TLS: Let\'s Encrypt certificate',
        'Document Root: /var/www/vibelinkgh.com/',
        'Gzip Compression: Enabled',
        'Caching: Browser caching headers configured'
    ]

    for spec in server_specs:
        doc.add_paragraph(spec, style='List Bullet')

    doc.add_heading('Nginx Configuration Example', level=3)
    nginx_config = '''
server {
    listen 80;
    listen [::]:80;
    server_name vibelinkgh.com www.vibelinkgh.com;
    return 301 https://$server_name$request_uri;
}

server {
    listen 443 ssl http2;
    listen [::]:443 ssl http2;
    server_name vibelinkgh.com www.vibelinkgh.com;

    root /var/www/vibelinkgh.com;
    index index.html;

    ssl_certificate /etc/letsencrypt/live/vibelinkgh.com/fullchain.pem;
    ssl_certificate_key /etc/letsencrypt/live/vibelinkgh.com/privkey.pem;

    # Gzip compression
    gzip on;
    gzip_types text/css application/javascript application/json image/svg+xml;
    gzip_comp_level 6;

    # Security headers
    add_header X-Frame-Options "SAMEORIGIN" always;
    add_header X-Content-Type-Options "nosniff" always;
    add_header X-XSS-Protection "1; mode=block" always;

    # Cache static assets
    location ~* \\.(js|css|png|jpg|jpeg|gif|svg|ico|woff|woff2)$ {
        expires 1y;
        add_header Cache-Control "public, immutable";
    }

    # SPA routing - redirect all to index.html
    location / {
        try_files $uri $uri/ /index.html;
    }
}
'''

    p = doc.add_paragraph(nginx_config)
    p.style = 'No Spacing'

    doc.add_heading('Domain Configuration', level=2)
    doc.add_paragraph('DNS records configuration for vibelinkgh.com:')

    dns_table = doc.add_table(rows=5, cols=4)
    dns_table.style = 'Light Grid Accent 1'

    dns_data = [
        ('Type', 'Name', 'Value', 'TTL'),
        ('A', '@', 'Server IP Address', '3600'),
        ('A', 'www', 'Server IP Address', '3600'),
        ('CNAME', 'www', 'vibelinkgh.com', '3600'),
        ('TXT', '@', 'SPF/DKIM records', '3600')
    ]

    for i, (record_type, name, value, ttl) in enumerate(dns_data):
        dns_table.cell(i, 0).text = record_type
        dns_table.cell(i, 1).text = name
        dns_table.cell(i, 2).text = value
        dns_table.cell(i, 3).text = ttl

    doc.add_heading('SSL Certificates', level=2)
    doc.add_paragraph(
        'SSL/TLS certificates are obtained from Let\'s Encrypt for free, automatically '
        'renewed HTTPS encryption.'
    )

    doc.add_heading('SSL Setup Steps', level=3)
    ssl_steps = [
        '1. Install Certbot: sudo apt install certbot python3-certbot-nginx',
        '2. Run Certbot: sudo certbot --nginx -d vibelinkgh.com -d www.vibelinkgh.com',
        '3. Follow prompts to configure certificate',
        '4. Test renewal: sudo certbot renew --dry-run',
        '5. Auto-renewal cron job: Certbot installs this automatically',
        '6. Verify HTTPS: Visit https://vibelinkgh.com',
        '7. Check SSL rating: Use SSL Labs test'
    ]

    for step in ssl_steps:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('Backup Strategy', level=2)
    doc.add_paragraph(
        'Comprehensive backup strategy ensures data protection and disaster recovery.'
    )

    doc.add_heading('Database Backups', level=3)
    db_backups = [
        'Supabase automatic daily backups (retained for 7 days on free tier)',
        'Manual backups before major migrations',
        'Export database to SQL file weekly',
        'Store backups in separate location (cloud storage)',
        'Test backup restoration quarterly',
        'Point-in-time recovery available (paid tier)'
    ]

    for backup in db_backups:
        doc.add_paragraph(backup, style='List Bullet')

    doc.add_heading('Code Backups', level=3)
    code_backups = [
        'GitHub repository serves as primary code backup',
        'Local git repository on development PC',
        'Tagged releases for version milestones',
        'Branch protection on main branch',
        'Regular commits with meaningful messages'
    ]

    for backup in code_backups:
        doc.add_paragraph(backup, style='List Bullet')

    doc.add_heading('File Storage Backups', level=3)
    storage_backups = [
        'Supabase Storage automatic replication',
        'Export uploaded files monthly',
        'Store backups in AWS S3 or similar',
        'Verify file integrity periodically'
    ]

    for backup in storage_backups:
        doc.add_paragraph(backup, style='List Bullet')

    doc.add_heading('Monitoring', level=2)
    doc.add_paragraph(
        'Proactive monitoring ensures system health and quick issue detection.'
    )

    doc.add_heading('Application Monitoring', level=3)
    app_monitoring = [
        'Supabase Dashboard: Database performance, API usage',
        'Error tracking: Console errors logged',
        'Performance metrics: Page load times, API response times',
        'User analytics: Google Analytics or similar',
        'Real-time user activity',
        'Uptime monitoring: UptimeRobot or Pingdom'
    ]

    for item in app_monitoring:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Server Monitoring', level=3)
    server_monitoring = [
        'CPU and memory usage',
        'Disk space utilization',
        'Network bandwidth',
        'Web server logs (access and error logs)',
        'SSL certificate expiration',
        'Security scan alerts'
    ]

    for item in server_monitoring:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Key Metrics to Monitor', level=3)
    metrics_table = doc.add_table(rows=9, cols=3)
    metrics_table.style = 'Light List Accent 1'

    metrics_data = [
        ('Metric', 'Target', 'Alert Threshold'),
        ('Uptime', '>99.5%', '<99%'),
        ('Page Load Time', '<2 seconds', '>3 seconds'),
        ('API Response Time', '<500ms', '>1000ms'),
        ('Error Rate', '<0.1%', '>1%'),
        ('Database Connections', 'Normal range', '>80% capacity'),
        ('Disk Space', '>20% free', '<10% free'),
        ('SSL Expiry', '>30 days', '<14 days')
    ]

    for i, (metric, target, threshold) in enumerate(metrics_data):
        metrics_table.cell(i, 0).text = metric
        metrics_table.cell(i, 1).text = target
        metrics_table.cell(i, 2).text = threshold

    doc.add_heading('Deployment Checklist', level=2)
    checklist = [
        '☐ Code tested in development environment',
        '☐ All tests passing',
        '☐ Environment variables configured',
        '☐ Build production bundle: npm run build',
        '☐ Test production build locally: npm run preview',
        '☐ Backup current production deployment',
        '☐ Backup database',
        '☐ Upload dist/ folder to server',
        '☐ Clear server cache',
        '☐ Verify website loads correctly',
        '☐ Test critical user flows (order, payment)',
        '☐ Check all integrations working',
        '☐ Monitor error logs for 1 hour',
        '☐ Update documentation',
        '☐ Notify team of deployment'
    ]

    for item in checklist:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

def add_security(doc):
    """Add security section"""
    doc.add_heading('SECURITY', level=1)

    doc.add_heading('Authentication', level=2)
    doc.add_paragraph(
        'Supabase Auth provides secure authentication with industry-standard practices.'
    )

    doc.add_heading('Authentication Methods', level=3)
    auth_methods = [
        'Email/Password authentication with secure password hashing (bcrypt)',
        'Email verification required for new accounts',
        'Password strength requirements enforced',
        'Rate limiting on login attempts',
        'Account lockout after failed attempts',
        'Session management with JWT tokens',
        'Token expiration and refresh mechanism',
        'Secure password reset flow'
    ]

    for method in auth_methods:
        doc.add_paragraph(method, style='List Bullet')

    doc.add_heading('Password Requirements', level=3)
    password_reqs = [
        'Minimum 8 characters',
        'At least one uppercase letter',
        'At least one lowercase letter',
        'At least one number',
        'At least one special character (recommended)',
        'Not a commonly used password',
        'Not the same as email address'
    ]

    for req in password_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('Session Management', level=3)
    session_mgmt = [
        'JWT tokens stored in httpOnly cookies',
        'Access token expires after 1 hour',
        'Refresh token used to obtain new access token',
        'Logout clears all session data',
        'Automatic logout on token expiration',
        'Session tracking in database',
        'Device/browser fingerprinting (optional)'
    ]

    for item in session_mgmt:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Authorization', level=2)
    doc.add_paragraph(
        'Role-based access control (RBAC) ensures users can only access authorized resources.'
    )

    doc.add_heading('User Roles', level=3)
    roles_table = doc.add_table(rows=4, cols=2)
    roles_table.style = 'Medium Grid 1 Accent 1'

    roles_data = [
        ('Role', 'Permissions'),
        ('Guest', 'View public pages, create account, place order'),
        ('Customer', 'View own orders, update profile, access customer portal, track orders'),
        ('Admin', 'Full access to all data, manage orders, manage customers, view analytics')
    ]

    for i, (role, permissions) in enumerate(roles_data):
        roles_table.cell(i, 0).text = role
        roles_table.cell(i, 1).text = permissions

    doc.add_heading('Row Level Security (RLS)', level=3)
    doc.add_paragraph(
        'Database-level security ensures users can only access data they own:'
    )

    rls_examples = [
        'Customers can SELECT only their own orders: auth.uid() = customer_id',
        'Customers can UPDATE only their own profile',
        'Admins bypass RLS with service role key (server-side only)',
        'Public tables (blog, portfolio) allow SELECT for all',
        'INSERT policies validate data before creation',
        'DELETE operations restricted to admin or soft delete only'
    ]

    for example in rls_examples:
        doc.add_paragraph(example, style='List Bullet')

    doc.add_heading('Data Protection', level=2)

    doc.add_heading('Data Encryption', level=3)
    encryption = [
        'All data in transit encrypted with TLS 1.3 (HTTPS)',
        'Database connections encrypted',
        'Supabase encrypts data at rest',
        'Sensitive fields (passwords) hashed with bcrypt',
        'Payment data never stored (handled by Paystack)',
        'API keys stored in environment variables, never in code'
    ]

    for item in encryption:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Data Privacy', level=3)
    privacy = [
        'GDPR compliance considerations',
        'Privacy policy clearly stated',
        'Cookie consent implementation',
        'User data deletion on request',
        'Data export capability',
        'Minimal data collection',
        'No selling of user data',
        'Third-party data sharing limited to necessary services'
    ]

    for item in privacy:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Input Validation', level=3)
    validation = [
        'Client-side validation with Zod schema',
        'Server-side validation in database constraints',
        'SQL injection prevention via parameterized queries',
        'XSS protection via React\'s built-in escaping',
        'CSRF protection with token validation',
        'File upload validation (type, size)',
        'Email format validation',
        'Phone number format validation'
    ]

    for item in validation:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Payment Security', level=2)
    doc.add_paragraph(
        'Payment processing security is critical for customer trust and regulatory compliance.'
    )

    doc.add_heading('PCI DSS Compliance', level=3)
    pci_compliance = [
        'No card data stored on VibeLink servers',
        'Paystack handles all card data (PCI DSS Level 1 certified)',
        'Payment form hosted on Paystack domain',
        'Tokenization of payment methods',
        'Secure payment gateway connection',
        'Regular security audits by Paystack'
    ]

    for item in pci_compliance:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Payment Flow Security', level=3)
    payment_security = [
        '1. Customer initiates payment',
        '2. Frontend calls Paystack API with public key',
        '3. Paystack popup opens (on Paystack domain)',
        '4. Customer enters payment details on Paystack',
        '5. Paystack processes payment',
        '6. Paystack sends webhook to backend with transaction reference',
        '7. Backend verifies transaction with Paystack API using secret key',
        '8. Order status updated only after verification',
        '9. Customer receives confirmation'
    ]

    for step in payment_security:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('API Security', level=3)
    api_security = [
        'API keys stored in environment variables',
        'Supabase RLS policies protect database access',
        'Rate limiting on API endpoints',
        'CORS configured to allow only vibelinkgh.com',
        'API authentication required for sensitive endpoints',
        'Webhook signature verification',
        'IP whitelisting for admin API access (optional)'
    ]

    for item in api_security:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Security Best Practices', level=2)
    best_practices = [
        'Regular security audits',
        'Keep dependencies up to date (npm audit)',
        'Monitor for security vulnerabilities',
        'Secure server configuration',
        'Regular backups',
        'Incident response plan',
        'Security training for team',
        'Logging and monitoring',
        'Penetration testing (annually)',
        'Bug bounty program (future consideration)'
    ]

    for practice in best_practices:
        doc.add_paragraph(practice, style='List Bullet')

    doc.add_page_break()

def add_maintenance_support(doc):
    """Add maintenance and support section"""
    doc.add_heading('MAINTENANCE & SUPPORT', level=1)

    doc.add_heading('Update Procedures', level=2)

    doc.add_heading('Application Updates', level=3)
    doc.add_paragraph('Process for updating the VibeLink Ghana application:')

    update_steps = [
        '1. Review update requirements and plan changes',
        '2. Create feature branch: git checkout -b feature/update-name',
        '3. Make code changes in development environment',
        '4. Test changes thoroughly locally',
        '5. Commit changes: git commit -m "Description"',
        '6. Push to GitHub: git push origin feature/update-name',
        '7. Create pull request for review',
        '8. Merge to main branch after approval',
        '9. Build production: npm run build',
        '10. Test production build: npm run preview',
        '11. Deploy to production server',
        '12. Monitor for issues post-deployment',
        '13. Update documentation as needed'
    ]

    for step in update_steps:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('Dependency Updates', level=3)
    dependency_updates = [
        'Check for updates weekly: npm outdated',
        'Read changelogs before updating',
        'Update patch versions: npm update',
        'Test after each update',
        'Update minor versions with caution',
        'Major version updates require thorough testing',
        'Security updates prioritized',
        'Lock file committed after updates'
    ]

    for item in dependency_updates:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Database Updates', level=3)
    db_update_process = [
        '1. Plan schema changes carefully',
        '2. Create migration file: supabase migration new migration_name',
        '3. Write SQL in migration file',
        '4. Test in development: supabase db reset',
        '5. Review for data loss or performance impact',
        '6. Backup production database',
        '7. Apply to production during low-traffic period',
        '8. Verify migration success',
        '9. Monitor database performance',
        '10. Rollback if issues occur'
    ]

    for step in db_update_process:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('Backup Procedures', level=2)

    doc.add_heading('Automated Backups', level=3)
    auto_backups = [
        'Supabase daily database backups (automatic)',
        'GitHub commits for code versioning',
        'Weekly database export to external storage',
        'Monthly full system backup',
        'Retention: Keep 7 daily, 4 weekly, 12 monthly backups'
    ]

    for backup in auto_backups:
        doc.add_paragraph(backup, style='List Bullet')

    doc.add_heading('Manual Backup Process', level=3)
    manual_backup = '''
1. Database Backup:
   - Login to Supabase Dashboard
   - Navigate to Database > Backups
   - Click "Create Backup"
   - Download SQL export
   - Store in secure location

2. Code Backup:
   - Ensure all changes committed to Git
   - git push origin main
   - Create tagged release: git tag v1.0.0
   - git push --tags

3. File Storage Backup:
   - Navigate to Supabase Storage
   - Download all buckets
   - Store in external backup location

4. Configuration Backup:
   - Export .env file (securely)
   - Document configuration changes
   - Store in encrypted location
'''

    p = doc.add_paragraph(manual_backup)
    p.style = 'No Spacing'

    doc.add_heading('Restore Procedures', level=3)
    restore_steps = [
        '1. Identify scope of restoration needed',
        '2. Retrieve appropriate backup files',
        '3. For database: Use Supabase restore or SQL import',
        '4. For code: git checkout <tag> or deploy from backup',
        '5. For files: Re-upload to Supabase Storage',
        '6. Verify data integrity after restore',
        '7. Test critical functionality',
        '8. Monitor for issues',
        '9. Document incident and restoration'
    ]

    for step in restore_steps:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('Troubleshooting', level=2)

    doc.add_heading('Common Issues and Solutions', level=3)

    # Issues table
    issues_table = doc.add_table(rows=9, cols=2)
    issues_table.style = 'Light Grid Accent 1'

    issues_data = [
        ('Issue', 'Solution'),
        ('Build fails', 'Check for TypeScript errors, run npm install, clear cache'),
        ('Page not loading', 'Check browser console, verify API connection, check RLS policies'),
        ('Payment not processing', 'Verify Paystack credentials, check webhook configuration'),
        ('Images not displaying', 'Check Supabase Storage permissions, verify file paths'),
        ('Slow page load', 'Optimize images, check bundle size, enable caching'),
        ('Database connection error', 'Verify Supabase credentials, check network, check quota'),
        ('Authentication issues', 'Clear cookies, check session expiry, verify auth configuration'),
        ('Deploy fails', 'Check server permissions, verify file paths, check disk space')
    ]

    for i, (issue, solution) in enumerate(issues_data):
        issues_table.cell(i, 0).text = issue
        issues_table.cell(i, 1).text = solution

    doc.add_heading('Debugging Process', level=3)
    debugging = [
        '1. Reproduce the issue',
        '2. Check browser console for errors',
        '3. Check network tab for failed requests',
        '4. Review server/database logs',
        '5. Verify configuration settings',
        '6. Test in isolation',
        '7. Search for similar issues online',
        '8. Consult documentation',
        '9. Ask for help if stuck',
        '10. Document solution for future reference'
    ]

    for step in debugging:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('Log Files', level=3)
    log_files = [
        'Browser Console: Frontend errors and warnings',
        'Network Tab: API requests and responses',
        'Supabase Logs: Database queries and errors',
        'Server Logs: Web server access and error logs (/var/log/nginx/)',
        'Build Logs: Vite build output',
        'Git Logs: Code change history (git log)'
    ]

    for log in log_files:
        doc.add_paragraph(log, style='List Bullet')

    doc.add_heading('Support Channels', level=2)

    support_channels = [
        ('Technical Documentation', 'This document and inline code comments'),
        ('GitHub Issues', 'Track bugs and feature requests'),
        ('Team Communication', 'Slack, email, or team chat'),
        ('Supabase Documentation', 'https://supabase.com/docs'),
        ('React Documentation', 'https://react.dev'),
        ('Community Forums', 'Stack Overflow, Reddit, Discord'),
        ('Paystack Support', 'https://paystack.com/support'),
        ('Vendor Support', 'Contact respective service providers')
    ]

    for channel, description in support_channels:
        p = doc.add_paragraph()
        p.add_run(f'{channel}: ').bold = True
        p.add_run(description)

    doc.add_heading('Maintenance Schedule', level=2)

    maintenance_table = doc.add_table(rows=6, cols=2)
    maintenance_table.style = 'Medium Grid 1 Accent 1'

    maintenance_data = [
        ('Frequency', 'Tasks'),
        ('Daily', 'Monitor error logs, check uptime, respond to support tickets'),
        ('Weekly', 'Review analytics, check for dependency updates, backup database'),
        ('Monthly', 'Security audit, performance review, update documentation'),
        ('Quarterly', 'Test backup restoration, penetration testing, user feedback review'),
        ('Annually', 'Comprehensive security audit, architecture review, renewal of SSL/domains')
    ]

    for i, (freq, tasks) in enumerate(maintenance_data):
        maintenance_table.cell(i, 0).text = freq
        maintenance_table.cell(i, 1).text = tasks

    doc.add_page_break()

def add_future_enhancements(doc):
    """Add future enhancements section"""
    doc.add_heading('FUTURE ENHANCEMENTS', level=1)

    doc.add_heading('Planned Features', level=2)

    doc.add_heading('Short-term (3-6 months)', level=3)
    short_term = [
        'Mobile app (React Native) for iOS and Android',
        'Advanced invitation editor with drag-and-drop',
        'Video invitations',
        'Live RSVP dashboard for event organizers',
        'Guest list CSV import/export',
        'SMS notifications for events',
        'Integration with calendar apps (Google Calendar, Outlook)',
        'Multi-language expansion (Twi, Ga, Ewe)',
        'Social media auto-posting',
        'Advanced analytics for admins'
    ]

    for feature in short_term:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('Medium-term (6-12 months)', level=3)
    medium_term = [
        'Event management features (seating charts, gift registries)',
        'Vendor marketplace (photographers, caterers, decorators)',
        'E-commerce for event supplies',
        'Ticketing system for paid events',
        'Live streaming integration for virtual events',
        'Photo sharing and album creation',
        'Guest check-in app with QR codes',
        'Automated reminder system',
        'Integration with event planning tools',
        'White-label solution for event planners'
    ]

    for feature in medium_term:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('Long-term (12+ months)', level=3)
    long_term = [
        'AI-powered design generator',
        'Voice invitations with AI text-to-speech',
        'Augmented Reality (AR) invitations',
        'Blockchain-based digital collectible invitations',
        'International expansion to other African countries',
        'Multi-vendor platform',
        'Franchise model for regional partners',
        'Corporate event management suite',
        'Integration with government event registries',
        'B2B SaaS offering for enterprises'
    ]

    for feature in long_term:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('Scalability Considerations', level=2)

    doc.add_heading('Performance Optimization', level=3)
    performance = [
        'Implement CDN for static assets (Cloudflare, AWS CloudFront)',
        'Image optimization and lazy loading',
        'Code splitting for faster initial load',
        'Caching strategies (Redis, service workers)',
        'Database query optimization',
        'API response caching',
        'Bundle size reduction',
        'Server-side rendering (SSR) consideration'
    ]

    for item in performance:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Infrastructure Scaling', level=3)
    infrastructure = [
        'Auto-scaling for traffic spikes',
        'Load balancing across multiple servers',
        'Database read replicas',
        'Horizontal scaling of application servers',
        'Queue system for background jobs (email, notifications)',
        'Microservices architecture for specific features',
        'Containerization with Docker',
        'Kubernetes for orchestration'
    ]

    for item in infrastructure:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Database Scaling', level=3)
    database_scaling = [
        'Database connection pooling',
        'Query optimization and indexing',
        'Partitioning large tables',
        'Archive old data to separate storage',
        'Consider NoSQL for specific use cases',
        'Implement caching layer (Redis)',
        'Database monitoring and performance tuning',
        'Upgrade to higher Supabase tier or self-hosted PostgreSQL'
    ]

    for item in database_scaling:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Business Growth Features', level=3)
    business_features = [
        'Subscription model for unlimited invitations',
        'Premium templates marketplace',
        'Partnership program with event venues',
        'Affiliate marketing program expansion',
        'Corporate packages with volume discounts',
        'API for third-party integrations',
        'Custom domain for premium customers',
        'Dedicated account manager for enterprise clients'
    ]

    for feature in business_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('Technology Upgrades', level=3)
    tech_upgrades = [
        'Migrate to Next.js for SSR/SSG capabilities',
        'Implement GraphQL for more efficient data fetching',
        'Add TypeScript strict mode',
        'Upgrade to latest React version with concurrent features',
        'Implement micro-frontends for team autonomy',
        'Add E2E testing with Playwright or Cypress',
        'CI/CD pipeline automation',
        'Feature flags for gradual rollouts'
    ]

    for upgrade in tech_upgrades:
        doc.add_paragraph(upgrade, style='List Bullet')

    doc.add_page_break()

def add_appendices(doc):
    """Add appendices section"""
    doc.add_heading('APPENDICES', level=1)

    # Appendix A: Package Dependencies
    doc.add_heading('Appendix A: Package Dependencies List', level=2)

    doc.add_heading('Production Dependencies', level=3)
    prod_deps = [
        '@dnd-kit/core: 6.3.1 - Drag and drop functionality',
        '@hookform/resolvers: 3.10.0 - Form validation resolvers',
        '@radix-ui/react-*: Various - Accessible UI components',
        '@supabase/supabase-js: 2.89.0 - Supabase client',
        '@tanstack/react-query: 5.83.0 - Data fetching and state management',
        '@tiptap/react: 3.14.0 - Rich text editor',
        'react: 18.3.1 - UI library',
        'react-dom: 18.3.1 - React DOM renderer',
        'react-router-dom: 6.30.1 - Client-side routing',
        'tailwindcss: 3.4.17 - CSS framework',
        'framer-motion: 12.23.26 - Animation library',
        'recharts: 2.15.4 - Chart library',
        'zod: 3.25.76 - Schema validation',
        'date-fns: 3.6.0 - Date utilities',
        'lucide-react: 0.462.0 - Icon library'
    ]

    for dep in prod_deps:
        doc.add_paragraph(dep, style='List Bullet')

    doc.add_heading('Development Dependencies', level=3)
    dev_deps = [
        'typescript: 5.8.3 - TypeScript compiler',
        'vite: 5.4.19 - Build tool',
        '@vitejs/plugin-react-swc: 3.11.0 - Vite React plugin',
        'eslint: 9.32.0 - Code linting',
        'autoprefixer: 10.4.21 - CSS autoprefixer',
        'postcss: 8.5.6 - CSS processor',
        '@types/react: 18.3.23 - React TypeScript definitions',
        '@types/node: 22.16.5 - Node.js TypeScript definitions'
    ]

    for dep in dev_deps:
        doc.add_paragraph(dep, style='List Bullet')

    # Appendix B: File Structure
    doc.add_heading('Appendix B: File Structure', level=2)

    file_structure = '''
vibelink/app/
├── public/                      # Static assets
│   ├── images/                  # Image files
│   ├── favicon.ico
│   └── robots.txt
├── src/
│   ├── components/              # React components
│   │   ├── admin/               # Admin-specific components
│   │   ├── auth/                # Authentication components
│   │   ├── customer/            # Customer portal components
│   │   ├── layout/              # Layout components (Header, Footer)
│   │   ├── order-form/          # Order form components
│   │   ├── payment/             # Payment components
│   │   ├── ui/                  # Base UI components (shadcn-ui)
│   │   ├── ChatWidget.tsx       # AI chatbot widget
│   │   ├── FloatingWhatsApp.tsx # WhatsApp button
│   │   └── SEO.tsx              # SEO component
│   ├── pages/                   # Page components (32 pages)
│   │   ├── Index.tsx            # Homepage
│   │   ├── About.tsx
│   │   ├── Services.tsx
│   │   ├── GetStarted.tsx       # Order form
│   │   ├── Admin.tsx            # Admin dashboard
│   │   └── ...
│   ├── integrations/
│   │   └── supabase/
│   │       ├── client.ts        # Supabase client
│   │       └── types.ts         # Database type definitions
│   ├── hooks/                   # Custom React hooks
│   ├── data/                    # Static data
│   ├── lib/                     # Utility functions
│   ├── assets/                  # Images, fonts
│   ├── App.tsx                  # Main app component
│   ├── main.tsx                 # Entry point
│   └── index.css                # Global styles
├── supabase/
│   └── migrations/              # 46 database migration files
├── dist/                        # Production build output
├── package.json                 # Dependencies
├── tsconfig.json                # TypeScript config
├── tailwind.config.ts           # Tailwind config
├── vite.config.ts               # Vite config
├── .env                         # Environment variables
└── README.md                    # Project documentation
'''

    p = doc.add_paragraph(file_structure)
    p.style = 'No Spacing'

    # Appendix C: Environment Variables
    doc.add_heading('Appendix C: Environment Variables', level=2)

    env_vars = [
        ('VITE_SUPABASE_URL', 'Supabase project URL', 'Required', 'Public'),
        ('VITE_SUPABASE_PUBLISHABLE_KEY', 'Supabase anon/public key', 'Required', 'Public'),
        ('VITE_SUPABASE_PROJECT_ID', 'Supabase project ID', 'Required', 'Public'),
        ('PAYSTACK_PUBLIC_KEY', 'Paystack public key', 'Required', 'Public'),
        ('PAYSTACK_SECRET_KEY', 'Paystack secret key', 'Required', 'Secret - Server only'),
        ('GEMINI_API_KEY', 'Google Gemini AI key', 'Required', 'Secret - Server only'),
        ('WHATSAPP_API_KEY', 'WhatsApp Business API key', 'Optional', 'Secret - Server only'),
        ('SMTP_HOST', 'Email SMTP host', 'Optional', 'Server only'),
        ('SMTP_USER', 'Email SMTP username', 'Optional', 'Server only'),
        ('SMTP_PASS', 'Email SMTP password', 'Optional', 'Secret - Server only')
    ]

    env_table = doc.add_table(rows=len(env_vars) + 1, cols=4)
    env_table.style = 'Light Grid Accent 1'
    env_table.cell(0, 0).text = 'Variable Name'
    env_table.cell(0, 1).text = 'Description'
    env_table.cell(0, 2).text = 'Required'
    env_table.cell(0, 3).text = 'Visibility'

    for i, (var, desc, req, vis) in enumerate(env_vars, 1):
        env_table.cell(i, 0).text = var
        env_table.cell(i, 1).text = desc
        env_table.cell(i, 2).text = req
        env_table.cell(i, 3).text = vis

    # Appendix D: Deployment Checklist
    doc.add_heading('Appendix D: Deployment Checklist', level=2)

    doc.add_heading('Pre-Deployment', level=3)
    pre_deploy = [
        '☐ All features tested in development',
        '☐ Code reviewed and approved',
        '☐ TypeScript compilation successful (no errors)',
        '☐ All environment variables configured',
        '☐ Database migrations tested',
        '☐ Security audit completed',
        '☐ Performance optimization done',
        '☐ SEO meta tags updated',
        '☐ Backup current production',
        '☐ Notify stakeholders of deployment'
    ]

    for item in pre_deploy:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('During Deployment', level=3)
    during_deploy = [
        '☐ Run production build: npm run build',
        '☐ Test build locally: npm run preview',
        '☐ Upload dist/ to server',
        '☐ Apply database migrations if any',
        '☐ Update environment variables on server',
        '☐ Clear server cache',
        '☐ Restart web server if needed',
        '☐ Verify SSL certificate valid'
    ]

    for item in during_deploy:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Post-Deployment', level=3)
    post_deploy = [
        '☐ Verify website loads (https://vibelinkgh.com)',
        '☐ Test user registration and login',
        '☐ Place test order',
        '☐ Test payment flow',
        '☐ Verify email notifications sent',
        '☐ Check admin dashboard accessible',
        '☐ Test mobile responsiveness',
        '☐ Monitor error logs for 1 hour',
        '☐ Check analytics tracking',
        '☐ Update documentation',
        '☐ Notify team of successful deployment'
    ]

    for item in post_deploy:
        doc.add_paragraph(item, style='List Bullet')

    # Appendix E: Glossary
    doc.add_heading('Appendix E: Glossary', level=2)

    glossary = [
        ('API', 'Application Programming Interface - Set of protocols for software communication'),
        ('CDN', 'Content Delivery Network - Distributed server network for fast content delivery'),
        ('CI/CD', 'Continuous Integration/Continuous Deployment - Automated software delivery'),
        ('CRUD', 'Create, Read, Update, Delete - Basic database operations'),
        ('DNS', 'Domain Name System - Translates domain names to IP addresses'),
        ('HMR', 'Hot Module Replacement - Live code updates without full page reload'),
        ('JWT', 'JSON Web Token - Secure token-based authentication'),
        ('ORM', 'Object-Relational Mapping - Database abstraction layer'),
        ('PWA', 'Progressive Web App - Web app with native-like features'),
        ('RLS', 'Row Level Security - Database-level access control'),
        ('RSVP', 'Répondez s\'il vous plaît - Please respond (invitation response)'),
        ('SaaS', 'Software as a Service - Cloud-based software delivery'),
        ('SEO', 'Search Engine Optimization - Improving search visibility'),
        ('SPA', 'Single Page Application - Web app that loads once and updates dynamically'),
        ('SQL', 'Structured Query Language - Database query language'),
        ('SSL/TLS', 'Secure Sockets Layer/Transport Layer Security - Encryption protocols'),
        ('UI/UX', 'User Interface/User Experience - Design and usability'),
        ('Webhook', 'HTTP callback - Automated message from one system to another')
    ]

    for term, definition in glossary:
        p = doc.add_paragraph()
        p.add_run(f'{term}: ').bold = True
        p.add_run(definition)

    doc.add_paragraph('\n\n')
    doc.add_paragraph('--- End of Document ---').alignment = WD_ALIGN_PARAGRAPH.CENTER

def setup_styles(doc):
    """Setup document styles"""
    # Create styles for consistent formatting
    styles = doc.styles

    # Normal style modifications
    style = styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Heading styles
    for i in range(1, 4):
        heading_style = styles[f'Heading {i}']
        heading_style.font.name = 'Cambria'
        heading_style.font.color.rgb = RGBColor(26, 35, 126)  # Navy blue

def main():
    """Main function to create the document"""
    print("Creating VibeLink Ghana Technical Documentation...")

    # Create new document
    doc = Document()

    # Setup styles
    setup_styles(doc)

    # Configure page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add all sections
    print("Adding cover page...")
    add_cover_page(doc)

    print("Adding table of contents...")
    add_table_of_contents(doc)

    print("Adding executive summary...")
    add_executive_summary(doc)

    print("Adding project overview...")
    add_project_overview(doc)

    print("Adding technical architecture...")
    add_technical_architecture(doc)

    print("Adding development workflow...")
    add_development_workflow(doc)

    print("Adding features and specifications...")
    add_features_specifications(doc)

    print("Adding user interface...")
    add_user_interface(doc)

    print("Adding database schema...")
    add_database_schema(doc)

    print("Adding API and integrations...")
    add_api_integrations(doc)

    print("Adding deployment and infrastructure...")
    add_deployment_infrastructure(doc)

    print("Adding security...")
    add_security(doc)

    print("Adding maintenance and support...")
    add_maintenance_support(doc)

    print("Adding future enhancements...")
    add_future_enhancements(doc)

    print("Adding appendices...")
    add_appendices(doc)

    # Save document
    output_path = r'C:\Users\CyberAware\OneDrive - Government of Ghana - CAGD\ZeroTrust\Visual Studio Code Workspace\vibelink\docs\VibeLink_Technical_Documentation.docx'
    print(f"Saving document to {output_path}...")
    doc.save(output_path)

    print("Documentation created successfully!")
    print(f"Location: {output_path}")
    print("\nNext steps:")
    print("1. Open the document in Microsoft Word")
    print("2. Generate table of contents (References > Table of Contents)")
    print("3. Add page numbers (Insert > Page Number)")
    print("4. Add headers (Insert > Header)")
    print("5. Create diagrams for architecture sections")
    print("6. Review and customize as needed")

if __name__ == '__main__':
    main()
